import csv
import io
import json
import markdown
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus.tableofcontents import SimpleIndex
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.units import mm
from docx import Document
from docx.shared import Inches
from flask import Blueprint, Response, abort, flash, redirect, render_template, request, url_for, jsonify, send_from_directory
from flask_login import current_user, login_required
from sqlalchemy import or_

from .. import db
from ..forms import EntryForm, SettingsForm
from ..services.email_service import email_service
from ..services.cache_service import cache_service
from ..models import User, Entry
from ..ai_service import ai_service


main_bp = Blueprint("main", __name__)


@main_bp.get("/")
def index():
    if current_user.is_authenticated:
        return redirect(url_for("main.dashboard"))
    return redirect(url_for("auth.login"))


@main_bp.get("/dashboard")
@login_required
def dashboard():
    q = (request.args.get("q") or "").strip()
    mood = (request.args.get("mood") or "").strip()
    category = (request.args.get("category") or "").strip()
    tag = (request.args.get("tag") or "").strip()
    favorite = (request.args.get("favorite") or "").strip()
    date_from = (request.args.get("date_from") or "").strip()
    date_to = (request.args.get("date_to") or "").strip()
    sentiment_filter = (request.args.get("sentiment") or "").strip()
    has_ai_insights = (request.args.get("ai_insights") or "").strip()
    page = request.args.get("page", default=1, type=int)
    per_page = request.args.get("per_page", default=10, type=int)
    sort_by = (request.args.get("sort_by") or "").strip()
    sort_order = (request.args.get("sort_order") or "").strip()

    query = Entry.query.filter_by(user_id=current_user.id)

    # Enhanced search with full-text capabilities
    if q:
        search_terms = q.split()
        search_conditions = []
        for term in search_terms:
            search_conditions.extend([
                Entry.title.ilike(f"%{term}%"),
                Entry.body.ilike(f"%{term}%"),
                Entry.tags.ilike(f"%{term}%"),
                Entry.category.ilike(f"%{term}%"),
                Entry.mood.ilike(f"%{term}%")
            ])
        query = query.filter(or_(*search_conditions))

    # Advanced filtering
    if mood:
        query = query.filter(Entry.mood == mood)
    
    if category:
        query = query.filter(Entry.category == category)
    
    if tag:
        query = query.filter(Entry.tags.ilike(f"%{tag}%"))
    
    if favorite == "1":
        query = query.filter(Entry.is_favorite.is_(True))
    
    # Date range filtering
    if date_from:
        try:
            from_date = datetime.strptime(date_from, '%Y-%m-%d')
            query = query.filter(Entry.created_at >= from_date)
        except ValueError:
            pass
    
    if date_to:
        try:
            to_date = datetime.strptime(date_to, '%Y-%m-%d')
            query = query.filter(Entry.created_at <= to_date)
        except ValueError:
            pass
    
    # Sentiment filtering
    if sentiment_filter:
        query = query.filter(Entry.sentiment == sentiment_filter)
    
    # AI insights filtering
    if has_ai_insights == "1":
        query = query.filter(Entry.ai_insights.isnot(None))
    
    # Advanced sorting
    if sort_by:
        if sort_by == "date":
            if sort_order == "asc":
                query = query.order_by(Entry.created_at.asc())
            else:
                query = query.order_by(Entry.created_at.desc())
        elif sort_by == "title":
            if sort_order == "asc":
                query = query.order_by(Entry.title.asc())
            else:
                query = query.order_by(Entry.title.desc())
        elif sort_by == "mood":
            if sort_order == "asc":
                query = query.order_by(Entry.mood.asc())
            else:
                query = query.order_by(Entry.mood.desc())
        elif sort_by == "updated":
            if sort_order == "asc":
                query = query.order_by(Entry.updated_at.asc())
            else:
                query = query.order_by(Entry.updated_at.desc())
        else:
            query = query.order_by(Entry.created_at.desc())
    else:
        query = query.order_by(Entry.created_at.desc())

    pagination = query.paginate(page=page, per_page=per_page, error_out=False)

    return render_template(
        "dashboard.html",
        entries=pagination.items,
        pagination=pagination,
        q=q,
        mood=mood,
        category=category,
        tag=tag,
        favorite=favorite,
        date_from=date_from,
        date_to=date_to,
        sentiment_filter=sentiment_filter,
        has_ai_insights=has_ai_insights,
        sort_by=sort_by,
        sort_order=sort_order,
        per_page=per_page,
    )


@main_bp.get("/entries/new")
@main_bp.post("/entries/new")
@login_required
def new_entry():
    form = EntryForm()
    suggestions = None
    sentiment_analysis = None
    
    if form.validate_on_submit():
        if form.get_suggestions.data:
            suggestions = ai_service.generate_entry_suggestions_with_search(
                mood=form.mood.data, 
                tags=form.tags.data,
                category=form.category.data
            )
            return render_template("entry_form.html", form=form, suggestions=suggestions)
        
        if form.analyze_sentiment.data and form.body.data:
            sentiment_analysis = ai_service.analyze_entry_sentiment_with_context(form.body.data)
            return render_template("entry_form.html", form=form, sentiment_analysis=sentiment_analysis)
        
        entry = Entry(
            title=form.title.data.strip(),
            body=form.body.data,
            mood=form.mood.data or None,
            category=form.category.data or None,
            tags=(form.tags.data or "").strip() or None,
            is_favorite=bool(form.is_favorite.data),
            user_id=current_user.id,
        )
        db.session.add(entry)
        db.session.commit()
        
        # Auto-analyze sentiment after saving
        try:
            sentiment_result = ai_service.analyze_entry_sentiment_with_context(entry.body)
            if sentiment_result.get("success"):
                entry.sentiment = sentiment_result.get("sentiment")
                entry.mood_score = sentiment_result.get("mood_score")
                entry.emotions = sentiment_result.get("emotions")
                entry.ai_insights = sentiment_result.get("insights")
                db.session.commit()
        except Exception as e:
            # Log error but don't fail the entry creation
            print(f"Enhanced sentiment analysis failed: {e}")
        
        flash("Entry saved.", "success")
        return redirect(url_for("main.dashboard"))

    return render_template("entry_form.html", form=form)


@main_bp.get("/entries/<int:entry_id>")
@login_required
def entry_detail(entry_id: int):
    entry = Entry.query.filter_by(id=entry_id, user_id=current_user.id).first()
    if not entry:
        abort(404)
    return render_template("entry_detail.html", entry=entry)


@main_bp.get("/entries/<int:entry_id>/edit")
@main_bp.post("/entries/<int:entry_id>/edit")
@login_required
def edit_entry(entry_id: int):
    entry = Entry.query.filter_by(id=entry_id, user_id=current_user.id).first()
    if not entry:
        abort(404)

    form = EntryForm(obj=entry)
    if request.method == "GET":
        form.mood.data = entry.mood or ""
        form.tags.data = entry.tags or ""
        form.is_favorite.data = bool(entry.is_favorite)

    if form.validate_on_submit():
        entry.title = form.title.data.strip()
        entry.body = form.body.data
        entry.mood = form.mood.data or None
        entry.tags = (form.tags.data or "").strip() or None
        entry.is_favorite = bool(form.is_favorite.data)
        entry.touch()
        db.session.commit()
        flash("Entry updated.", "success")
        return redirect(url_for("main.entry_detail", entry_id=entry.id))

    return render_template("entry_form.html", form=form, entry=entry)


@main_bp.post("/entries/<int:entry_id>/delete")
@login_required
def delete_entry(entry_id: int):
    entry = Entry.query.filter_by(id=entry_id, user_id=current_user.id).first()
    if not entry:
        abort(404)
    db.session.delete(entry)
    db.session.commit()
    flash("Entry deleted.", "success")
    return redirect(url_for("main.dashboard"))


@main_bp.post("/entries/<int:entry_id>/favorite")
@login_required
def toggle_favorite(entry_id: int):
    entry = Entry.query.filter_by(id=entry_id, user_id=current_user.id).first()
    if not entry:
        abort(404)
    entry.is_favorite = not bool(entry.is_favorite)
    entry.touch()
    db.session.commit()
    return redirect(request.referrer or url_for("main.dashboard"))


@main_bp.get("/mood-analytics")
@login_required
def mood_analytics():
    # Get user's entries with mood data
    entries = Entry.query.filter_by(user_id=current_user.id).order_by(Entry.created_at.desc()).limit(50).all()
    
    # Prepare data for AI analysis
    entries_data = []
    for entry in entries:
        if entry.mood_score is not None:
            entries_data.append({
                'mood_score': entry.mood_score,
                'created_at': entry.created_at
            })
    
    # Get mood trends
    mood_trends = ai_service.get_mood_trends_with_research(entries_data)
    
    return render_template("mood_analytics.html", entries=entries, mood_trends=mood_trends)


@main_bp.get("/wellness")
@login_required
def wellness():
    entries = (
        Entry.query.filter_by(user_id=current_user.id)
        .order_by(Entry.created_at.desc())
        .limit(10)
        .all()
    )
    
    if not entries:
        flash("Write some entries first to get wellness insights.", "info")
        return redirect(url_for("main.dashboard"))
    
    entries_text = [entry.body for entry in entries]
    insights = ai_service.generate_wellness_insights_with_research(entries_text)
    
    mood_counts = {}
    for entry in entries:
        if entry.mood:
            mood_counts[entry.mood] = mood_counts.get(entry.mood, 0) + 1
    
    return render_template("wellness.html", 
                         insights=insights, 
                         mood_counts=mood_counts,
                         recent_entries=entries[:5])


@main_bp.get("/advanced-analytics")
@login_required
def advanced_analytics():
    """Display advanced analytics dashboard with writing patterns"""
    # Get user's entries with analytics data
    entries = Entry.query.filter_by(user_id=current_user.id).order_by(Entry.created_at.desc()).all()
    
    # Calculate writing patterns
    analytics_data = {
        'total_entries': len(entries),
        'total_words': sum(len(entry.body.split()) for entry in entries),
        'avg_words_per_entry': 0,
        'writing_streak': 0,
        'most_productive_day': None,
        'category_distribution': {},
        'mood_distribution': {},
        'monthly_writing': {},
        'tag_frequency': {},
        'writing_times': {}
    }
    
    if entries:
        # Average words per entry
        analytics_data['avg_words_per_entry'] = analytics_data['total_words'] / len(entries)
        
        # Writing streak (consecutive days)
        dates = [entry.created_at.date() for entry in entries]
        analytics_data['writing_streak'] = calculate_writing_streak(dates)
        
        # Most productive day of week
        day_counts = {}
        for date in dates:
            day_name = date.strftime('%A')
            day_counts[day_name] = day_counts.get(day_name, 0) + 1
        if day_counts:
            analytics_data['most_productive_day'] = max(day_counts, key=day_counts.get)
        
        # Category distribution
        for entry in entries:
            category = entry.category or 'Uncategorized'
            analytics_data['category_distribution'][category] = analytics_data['category_distribution'].get(category, 0) + 1
        
        # Mood distribution
        for entry in entries:
            mood = entry.mood or 'Neutral'
            analytics_data['mood_distribution'][mood] = analytics_data['mood_distribution'].get(mood, 0) + 1
        
        # Monthly writing patterns
        for entry in entries:
            month = entry.created_at.strftime('%Y-%m')
            analytics_data['monthly_writing'][month] = analytics_data['monthly_writing'].get(month, 0) + 1
        
        # Tag frequency
        for entry in entries:
            if entry.tags:
                tags = [tag.strip() for tag in entry.tags.split(',')]
                for tag in tags:
                    if tag:
                        analytics_data['tag_frequency'][tag] = analytics_data['tag_frequency'].get(tag, 0) + 1
        
        # Sort tag frequency by count and get top 10
        analytics_data['tag_frequency'] = dict(sorted(analytics_data['tag_frequency'].items(), key=lambda x: x[1], reverse=True)[:10])
        
        # Writing time patterns (hour of day)
        for entry in entries:
            hour = entry.created_at.hour
            time_period = get_time_period(hour)
            analytics_data['writing_times'][time_period] = analytics_data['writing_times'].get(time_period, 0) + 1
    
    return render_template("advanced_analytics.html", analytics=analytics_data, entries=entries[:30])


def calculate_writing_streak(dates):
    """Calculate current writing streak"""
    if not dates:
        return 0
    
    # Get unique dates and sort
    unique_dates = sorted(list(set(dates)), reverse=True)
    
    streak = 1  # Current day counts as streak of 1
    today = datetime.now().date()
    
    if unique_dates[0] != today:
        return 0  # No entry today, streak is 0
    
    for i in range(1, len(unique_dates)):
        prev_date = unique_dates[i-1]
        curr_date = unique_dates[i]
        
        # Check if current date is exactly one day before previous
        if (prev_date - curr_date).days == 1:
            streak += 1
        else:
            break
    
    return streak


def get_time_period(hour):
    """Convert hour to time period"""
    if 5 <= hour < 12:
        return "Morning (5AM-12PM)"
    elif 12 <= hour < 17:
        return "Afternoon (12PM-5PM)"
    elif 17 <= hour < 21:
        return "Evening (5PM-9PM)"
    else:
        return "Night (9PM-5AM)"


@main_bp.get("/export")
@login_required
def export_page():
    """Display export options page"""
    return render_template("export.html")


@main_bp.get("/export/<string:fmt>")
@login_required
def export_entries(fmt: str):
    entries = (
        Entry.query.filter_by(user_id=current_user.id)
        .order_by(Entry.created_at.desc())
        .all()
    )

    if fmt == "json":
        payload = [
            {
                "id": e.id,
                "title": e.title,
                "body": e.body,
                "mood": e.mood,
                "tags": e.tags,
                "category": e.category,
                "is_favorite": bool(e.is_favorite),
                "sentiment": e.sentiment,
                "mood_score": e.mood_score,
                "emotions": json.loads(e.emotions) if e.emotions else [],
                "ai_insights": e.ai_insights,
                "created_at": e.created_at.isoformat(),
                "updated_at": e.updated_at.isoformat() if e.updated_at else None,
            }
            for e in entries
        ]
        data = json.dumps(payload, ensure_ascii=False, indent=2)
        return Response(
            data,
            mimetype="application/json",
            headers={"Content-Disposition": "attachment; filename=diary-export.json"},
        )

    if fmt == "txt":
        out = io.StringIO()
        for e in entries:
            out.write(f"{e.created_at:%Y-%m-%d %H:%M} - {e.title}\n")
            if e.mood:
                out.write(f"Mood: {e.mood}\n")
            if e.tags:
                out.write(f"Tags: {e.tags}\n")
            out.write(e.body)
            out.write("\n\n" + ("-" * 40) + "\n\n")
        return Response(
            out.getvalue(),
            mimetype="text/plain; charset=utf-8",
            headers={"Content-Disposition": "attachment; filename=diary-export.txt"},
        )

    if fmt == "csv":
        out = io.StringIO(newline="")
        writer = csv.writer(out)
        writer.writerow(["id", "title", "mood", "tags", "is_favorite", "created_at", "updated_at", "body"])
        for e in entries:
            writer.writerow(
                [
                    e.id,
                    e.title,
                    e.mood or "",
                    e.tags or "",
                    int(bool(e.is_favorite)),
                    e.created_at.isoformat(),
                    e.updated_at.isoformat() if e.updated_at else "",
                    e.body,
                ]
            )
        return Response(
            out.getvalue(),
            mimetype="text/csv; charset=utf-8",
            headers={"Content-Disposition": "attachment; filename=diary-export.csv"},
        )

    if fmt == "docx":
        buffer = io.BytesIO()
        doc = Document()
        
        # Add title
        title = doc.add_heading('My Diary Export', 0)
        title.alignment = 1  # Center
        
        for i, e in enumerate(entries, 1):
            # Entry title
            doc.add_heading(f'Entry {i}: {e.title}', level=1)
            
            # Entry metadata
            p = doc.add_paragraph()
            p.add_run('Date: ').bold = True
            p.add_run(f'{e.created_at.strftime("%Y-%m-%d %H:%M")}\n')
            
            if e.mood:
                p.add_run('Mood: ').bold = True
                p.add_run(f'{e.mood.title()}\n')
            
            if e.category:
                p.add_run('Category: ').bold = True
                p.add_run(f'{e.category.title()}\n')
            
            if e.tags:
                p.add_run('Tags: ').bold = True
                p.add_run(f'{e.tags}\n')
            
            if e.sentiment:
                p.add_run('Sentiment: ').bold = True
                p.add_run(f'{e.sentiment.title()}\n')
            
            if e.ai_insights:
                p.add_run('AI Insights: ').bold = True
                p.add_run(f'{e.ai_insights[:200]}...\n')
            
            # Entry body
            doc.add_paragraph(e.body)
            
            # Separator
            if i < len(entries):
                doc.add_page_break()
        
        doc.save(buffer)
        buffer.seek(0)
        return Response(
            buffer.getvalue(),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=diary-export.docx"},
        )

    if fmt == "pdf":
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Get theme from request or default
        theme = request.args.get('theme', 'default')
        
        # Define theme colors
        if theme == 'dark':
            title_color = colors.HexColor('#2d3748')
            text_color = colors.HexColor('#e9ecef')
            bg_color = colors.HexColor('#212529')
            border_color = colors.HexColor('#495057')
        elif theme == 'nature':
            title_color = colors.HexColor('#2e7d32')
            text_color = colors.HexColor('#383e3c')
            bg_color = colors.HexColor('#f8f9fa')
            border_color = colors.HexColor('#28a745')
        elif theme == 'ocean':
            title_color = colors.HexColor('#006994')
            text_color = colors.HexColor('#343a40')
            bg_color = colors.HexColor('#e3f2fd')
            border_color = colors.HexColor('#004085')
        else:  # default
            title_color = colors.HexColor('#007bff')
            text_color = colors.HexColor('#212529')
            bg_color = colors.HexColor('#f8f9fa')
            border_color = colors.HexColor('#dee2e6')
        
        # Title page
        title_style = styles['Title']
        title_style.textColor = title_color
        story.append(Paragraph("My Diary Export", title_style))
        story.append(Spacer(1, 20))
        
        # Add metadata
        meta_style = styles['Normal']
        meta_style.textColor = text_color
        meta_style.backColor = bg_color
        story.append(Paragraph(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M')}", meta_style))
        story.append(Paragraph(f"Total Entries: {len(entries)}", meta_style))
        story.append(Spacer(1, 20))
        
        for i, e in enumerate(entries, 1):
            # Entry title with theme
            title_style = styles['Heading1']
            title_style.textColor = title_color
            story.append(Paragraph(f"Entry {i}: {e.title}", title_style))
            story.append(Spacer(1, 12))
            
            # Entry metadata box
            meta_text = f"""
            <font color="{text_color.hex if hasattr(text_color, 'hex') else '#000000'}">
                <b>Date:</b> {e.created_at.strftime('%Y-%m-%d %H:%M')}<br/>
                {f'<b>Mood:</b> {e.mood.title()}<br/>' if e.mood else ''}
                {f'<b>Category:</b> {e.category.title()}<br/>' if e.category else ''}
                {f'<b>Sentiment:</b> {e.sentiment.title()}<br/>' if e.sentiment else ''}
                {f'<b>Score:</b> {e.mood_score:.1f}/1.0<br/>' if e.mood_score else ''}
            </font>
            """
            story.append(Paragraph(meta_text, meta_style))
            story.append(Spacer(1, 12))
            
            # Entry body
            body_style = styles['BodyText']
            body_style.textColor = text_color
            body_style.backColor = bg_color
            body_style.borderColor = border_color
            body_style.borderWidth = 1
            body_style.borderRadius = 5
            body_style.leftIndent = 10
            body_style.rightIndent = 10
            body_style.topPadding = 10
            body_style.bottomPadding = 10
            
            # Convert markdown to HTML for better formatting
            body_html = markdown.markdown(e.body)
            story.append(Paragraph(body_html, body_style))
            
            # AI Insights if available
            if e.ai_insights:
                insights_style = styles['Normal']
                insights_style.textColor = text_color
                insights_style.backColor = colors.HexColor('#e9ecef')
                insights_style.borderWidth = 1
                insights_style.borderColor = border_color
                insights_style.borderRadius = 3
                insights_style.leftIndent = 10
                insights_style.rightIndent = 10
                insights_style.topPadding = 8
                insights_style.bottomPadding = 8
                
                story.append(Spacer(1, 12))
                story.append(Paragraph("<b>AI Insights:</b>", insights_style))
                story.append(Paragraph(e.ai_insights[:300], insights_style))
            
            # Page break except for last entry
            if i < len(entries):
                story.append(Spacer(1, 20))
                story.append(Paragraph("<hr/>", meta_style))
                story.append(Spacer(1, 10))
        
        doc.build(story)
        buffer.seek(0)
        return Response(
            buffer.getvalue(),
            mimetype="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=diary-export-{theme}.pdf"},
        )

    if fmt == "md":
        out = io.StringIO()
        out.write("# My Diary Export\n\n")
        for e in entries:
            out.write(f"## {e.title}\n\n")
            out.write(f"**Date:** {e.created_at.strftime('%Y-%m-%d %H:%M')}\n")
            if e.mood:
                out.write(f"**Mood:** {e.mood.title()}\n")
            if e.tags:
                out.write(f"**Tags:** {e.tags}\n")
            out.write(f"\n{e.body}\n\n---\n\n")
        return Response(
            out.getvalue(),
            mimetype="text/markdown",
            headers={"Content-Disposition": "attachment; filename=diary-export.md"},
        )

    abort(404)


@main_bp.get("/settings")
@main_bp.post("/settings")
@login_required
def settings():
    form = SettingsForm()
    if request.method == "GET":
        form.theme.data = getattr(current_user, "theme", "dark")

    if form.validate_on_submit():
        current_user.theme = form.theme.data
        db.session.commit()
        flash("Settings saved.", "success")
        return redirect(url_for("main.settings"))

    return render_template("settings.html", form=form)


@main_bp.route("/ads.txt")
def ads_txt():
    """Serve ads.txt file for AdSense verification"""
    try:
        return send_from_directory('static', 'ads.txt', mimetype='text/plain')
    except:
        # Fallback if file not found
        return "google.com, pub-2396098605485959, DIRECT, f08c47fec0942fa0", 200, {'Content-Type': 'text/plain'}


# Email notification routes
@main_bp.post("/email-preferences")
@login_required
def update_email_preferences():
    """Update email notification preferences"""
    daily_reminders = request.form.get('daily_reminders') == 'on'
    weekly_summary = request.form.get('weekly_summary') == 'on'
    monthly_insights = request.form.get('monthly_insights') == 'on'
    
    current_user.email_daily_reminders = daily_reminders
    current_user.email_weekly_summary = weekly_summary
    current_user.email_monthly_insights = monthly_insights
    db.session.commit()
    
    flash("Email preferences updated successfully!", "success")
    return redirect(url_for('main.settings'))


@main_bp.get("/test-email")
@login_required
def test_email():
    """Send test email to user"""
    try:
        success = email_service.send_daily_reminder(
            current_user.email,
            current_user.email.split('@')[0].title()
        )
        if success:
            flash("Test email sent successfully!", "success")
        else:
            flash("Failed to send test email. Please check your email configuration.", "error")
    except Exception as e:
        flash(f"Error sending test email: {str(e)}", "error")
    
    return redirect(url_for('main.settings'))
